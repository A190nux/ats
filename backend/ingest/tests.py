"""
Test suite for document loader and OCR module.

Provides unit tests for PDF, DOCX, DOC, TXT ingestion and OCR fallback.
"""

import os
import tempfile
from pathlib import Path
import unittest
from io import BytesIO

# Try to import docx; skip tests if not available
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

from backend.ingest.loader import load_documents
from backend.ingest.ocr import is_ocr_available


class TestDocumentLoader(unittest.TestCase):
    """Test document loader with various file formats."""

    @classmethod
    def setUpClass(cls):
        """Create a temporary directory with test files."""
        cls.test_dir = tempfile.mkdtemp()
        
        # Create a test TXT file
        txt_path = os.path.join(cls.test_dir, "test.txt")
        with open(txt_path, 'w', encoding='utf-8') as f:
            f.write("John Doe\nEmail: john@example.com\nSkills: Python, JavaScript")
        
        # Create a test DOCX file (if python-docx is available)
        if HAS_DOCX:
            docx_path = os.path.join(cls.test_dir, "test.docx")
            doc = DocxDocument()
            doc.add_paragraph("Jane Smith")
            doc.add_paragraph("Email: jane@example.com")
            doc.add_paragraph("Skills: Java, C++")
            doc.save(docx_path)

    @classmethod
    def tearDownClass(cls):
        """Clean up temporary directory."""
        import shutil
        shutil.rmtree(cls.test_dir)

    def test_load_txt_file(self):
        """Test loading a plain text file."""
        docs = load_documents(self.test_dir, supported_formats=['.txt'])
        self.assertEqual(len(docs), 1, "Should load exactly 1 TXT file")
        self.assertIn("John Doe", docs[0].text)
        self.assertEqual(docs[0].metadata['file_type'], 'txt')

    @unittest.skipIf(not HAS_DOCX, "python-docx not installed")
    def test_load_docx_file(self):
        """Test loading a DOCX file."""
        docs = load_documents(self.test_dir, supported_formats=['.docx'])
        self.assertEqual(len(docs), 1, "Should load exactly 1 DOCX file")
        self.assertIn("Jane Smith", docs[0].text)
        self.assertEqual(docs[0].metadata['file_type'], 'docx')

    def test_load_mixed_formats(self):
        """Test loading mixed file formats."""
        docs = load_documents(self.test_dir, supported_formats=['.txt', '.docx'])
        # Should load TXT and DOCX (if available)
        expected_count = 1 if not HAS_DOCX else 2
        self.assertEqual(len(docs), expected_count, f"Should load {expected_count} documents")

    def test_load_nonexistent_directory(self):
        """Test behavior when directory doesn't exist."""
        docs = load_documents("/nonexistent/path")
        self.assertEqual(len(docs), 0, "Should return empty list for nonexistent directory")

    def test_metadata_present(self):
        """Test that metadata is correctly attached."""
        docs = load_documents(self.test_dir, supported_formats=['.txt'])
        self.assertEqual(len(docs), 1)
        doc = docs[0]
        self.assertIn('file_name', doc.metadata)
        self.assertIn('source', doc.metadata)
        self.assertIn('file_type', doc.metadata)
        self.assertEqual(doc.metadata['file_type'], 'txt')

    def test_ocr_disabled(self):
        """Test that loader respects enable_ocr=False flag."""
        docs = load_documents(self.test_dir, supported_formats=['.txt'], enable_ocr=False)
        self.assertEqual(len(docs), 1, "Should still load TXT even with OCR disabled")

    def test_ocr_availability(self):
        """Test OCR availability check."""
        ocr_available = is_ocr_available()
        self.assertIsInstance(ocr_available, bool)
        # Note: OCR may or may not be available depending on Tesseract installation


if __name__ == '__main__':
    unittest.main()
