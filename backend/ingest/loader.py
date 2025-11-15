"""
Robust document loader supporting PDF, DOCX, DOC, TXT, and image files.

Handles:
- PDF files (via pypdf, with OCR fallback for scanned PDFs)
- DOCX files (via python-docx)
- Legacy DOC files (via mammoth, fallback to striprtf)
- TXT files (plain text)
- Image files (JPG, PNG) — parsed via OCR (pytesseract)

Returns a list of llama_index.core.Document objects with metadata.
"""

import os
import logging
from pathlib import Path
from typing import List, Optional
import pypdf
from llama_index.core import Document

from backend.ingest.ocr import ocr_pdf, ocr_image, _detect_scanned_pdf

logger = logging.getLogger(__name__)

# Optional dependencies with graceful fallback
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    logger.warning("python-docx not installed. DOCX files will be skipped.")

try:
    import mammoth
    HAS_MAMMOTH = True
except ImportError:
    HAS_MAMMOTH = False
    logger.warning("mammoth not installed. Legacy DOC files may fail; will fallback to striprtf.")

try:
    import striprtf.rtf as rtf
    HAS_STRIPRTF = True
except ImportError:
    HAS_STRIPRTF = False
    logger.warning("striprtf not installed. RTF/DOC fallback will be skipped.")


def _load_pdf(filepath: str, enable_ocr: bool = True) -> Optional[str]:
    """Load and extract text from a PDF file.
    
    Falls back to OCR if the PDF is scanned (image-based).
    
    Args:
        filepath: Path to PDF file.
        enable_ocr: If True, use OCR for scanned PDFs.
    
    Returns:
        Extracted text or None if failed.
    """
    try:
        with open(filepath, 'rb') as file:
            pdf = pypdf.PdfReader(file)
            text = '\n'.join(page.extract_text() for page in pdf.pages if page.extract_text())
            
            # If we got minimal text and OCR is enabled, try OCR
            if enable_ocr and len(text.strip()) < 100 and _detect_scanned_pdf(filepath):
                logger.info(f"PDF appears scanned, attempting OCR: {filepath}")
                ocr_text = ocr_pdf(filepath)
                if ocr_text:
                    return ocr_text
            
            if text.strip():
                return text
            else:
                logger.warning(f"PDF {filepath} produced no extractable text (may be scanned; OCR disabled or unavailable).")
                return None
    except Exception as e:
        logger.error(f"Failed to load PDF {filepath}: {e}")
        return None


def _load_docx(filepath: str) -> Optional[str]:
    """Load and extract text from a DOCX file."""
    if not HAS_DOCX:
        logger.warning(f"Skipping DOCX {filepath}: python-docx not installed.")
        return None

    try:
        doc = DocxDocument(filepath)
        text = '\n'.join(para.text for para in doc.paragraphs if para.text.strip())
        if text.strip():
            return text
        else:
            logger.warning(f"DOCX {filepath} produced no extractable text.")
            return None
    except Exception as e:
        logger.error(f"Failed to load DOCX {filepath}: {e}")
        return None


def _load_doc_mammoth(filepath: str) -> Optional[str]:
    """Try loading legacy DOC file using mammoth."""
    if not HAS_MAMMOTH:
        logger.debug(f"mammoth not available for {filepath}.")
        return None

    try:
        result = mammoth.convert_file(filepath)
        text = result.value.strip()
        if text:
            return text
        else:
            logger.warning(f"DOC {filepath} produced no text via mammoth.")
            return None
    except Exception as e:
        logger.debug(f"mammoth failed for {filepath}: {e}")
        return None


def _load_doc_striprtf(filepath: str) -> Optional[str]:
    """Fallback: try loading DOC as RTF using striprtf."""
    if not HAS_STRIPRTF:
        logger.debug(f"striprtf not available for {filepath}.")
        return None

    try:
        with open(filepath, 'rb') as f:
            text = rtf.loads(f.read().decode('utf-8', errors='ignore')).strip()
            if text:
                return text
            else:
                logger.warning(f"DOC {filepath} produced no text via striprtf.")
                return None
    except Exception as e:
        logger.debug(f"striprtf failed for {filepath}: {e}")
        return None


def _load_doc(filepath: str) -> Optional[str]:
    """Load legacy DOC file with fallback chain: mammoth -> striprtf."""
    text = _load_doc_mammoth(filepath)
    if text:
        return text
    text = _load_doc_striprtf(filepath)
    if text:
        return text
    logger.error(f"Could not load DOC {filepath} with any available method.")
    return None


def _load_txt(filepath: str) -> Optional[str]:
    """Load plain text file."""
    try:
        with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
            text = f.read().strip()
            if text:
                return text
            else:
                logger.warning(f"TXT {filepath} is empty.")
                return None
    except Exception as e:
        logger.error(f"Failed to load TXT {filepath}: {e}")
        return None


def _load_image(filepath: str) -> Optional[str]:
    """Load image file (JPG, PNG, TIFF, etc.) via OCR."""
    return ocr_image(filepath)


def load_documents(directory: str, supported_formats: Optional[List[str]] = None, enable_ocr: bool = True) -> List[Document]:
    """
    Load documents from a directory, supporting multiple formats.

    Args:
        directory: Path to directory containing documents.
        supported_formats: List of file extensions to load (e.g., ['.pdf', '.docx', '.txt', '.jpg']).
                          If None, defaults to ['.pdf', '.docx', '.doc', '.txt', '.jpg', '.png', '.jpeg', '.tiff'].
        enable_ocr: If True, use OCR for scanned PDFs and image files.

    Returns:
        List of llama_index.core.Document objects with metadata.
    """
    if supported_formats is None:
        supported_formats = ['.pdf', '.docx', '.doc', '.txt', '.jpg', '.png', '.jpeg', '.tiff']

    # Normalize to lowercase for comparison
    supported_formats = [fmt.lower() for fmt in supported_formats]

    documents = []
    
    if not os.path.exists(directory):
        logger.error(f"Directory {directory} does not exist.")
        return documents

    for filename in os.listdir(directory):
        filepath = os.path.join(directory, filename)

        # Skip directories and hidden files
        if os.path.isdir(filepath) or filename.startswith('.'):
            continue

        file_ext = Path(filename).suffix.lower()

        # Skip unsupported formats
        if file_ext not in supported_formats:
            continue

        logger.info(f"Loading {filename}...")
        text = None

        if file_ext == '.pdf':
            text = _load_pdf(filepath, enable_ocr=enable_ocr)
        elif file_ext == '.docx':
            text = _load_docx(filepath)
        elif file_ext == '.doc':
            text = _load_doc(filepath)
        elif file_ext == '.txt':
            text = _load_txt(filepath)
        elif file_ext in ('.jpg', '.png', '.jpeg', '.tiff'):
            if enable_ocr:
                text = _load_image(filepath)
            else:
                logger.warning(f"Skipping image {filename} (OCR disabled)")

        if text:
            doc = Document(
                text=text,
                metadata={
                    "file_name": filename,
                    "source": filepath,
                    "file_type": file_ext[1:] if file_ext.startswith('.') else file_ext,
                }
            )
            documents.append(doc)
            logger.info(f"  ✓ Loaded {filename} ({len(text)} chars)")
        else:
            logger.warning(f"  ✗ Failed or produced empty text: {filename}")

    logger.info(f"Total documents loaded: {len(documents)}")
    return documents
